#!/usr/bin/env python3
"""把研究計畫書 markdown 轉成 Word (.docx)。只支援本文件用到的語法子集。"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

CJK_FONT = "PMingLiU"  # 繁中襯線字型；Word 端有則用之，否則回退
LATIN_FONT = "Calibri"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = LATIN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_formatted(paragraph, text, base_size=11, base_bold=False, base_color=None):
    """處理 **bold** 與 `code` 行內標記。"""
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            set_run_font(r, base_size, True, base_color)
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            rpr = r._element.get_or_add_rPr()
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rfonts.set(qn("w:eastAsia"), CJK_FONT)
            rpr.append(rfonts)
            r.font.size = Pt(base_size - 0.5)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            r = paragraph.add_run(part)
            set_run_font(r, base_size, base_bold, base_color)


def main(src, dst):
    with open(src, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    # 預設樣式字型
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 程式碼區塊
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳過結尾 ```
            for cl in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                r = p.add_run(cl if cl else " ")
                r.font.name = "Consolas"
                r.font.size = Pt(10)
                rpr = r._element.get_or_add_rPr()
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rfonts.set(qn("w:eastAsia"), CJK_FONT)
                rpr.append(rfonts)
            continue

        # 表格
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s:\-|]+\|?$", lines[i + 1].strip()):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # 跳過表頭與分隔列
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].text = ""
                add_formatted(cell.paragraphs[0], h, base_size=10, base_bold=True)
            for row in rows:
                cells = table.add_row().cells
                for j in range(len(header)):
                    txt = row[j] if j < len(row) else ""
                    cells[j].paragraphs[0].text = ""
                    add_formatted(cells[j].paragraphs[0], txt, base_size=10)
            doc.add_paragraph()
            continue

        # 水平線
        if stripped == "---":
            doc.add_paragraph().add_run().add_break()
            i += 1
            continue

        # 標題
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph()
            sizes = {1: 20, 2: 16, 3: 13, 4: 12, 5: 11, 6: 11}
            add_formatted(p, text, base_size=sizes.get(level, 11), base_bold=True,
                          base_color=RGBColor(0x1F, 0x49, 0x7D))
            p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            quote_text = re.sub(r"^>\s?", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_formatted(p, quote_text, base_color=RGBColor(0x55, 0x55, 0x55))
            i += 1
            continue

        # 有序清單
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Pt(18 + indent * 6)
            add_formatted(p, m.group(3))
            i += 1
            continue

        # 無序清單
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Pt(18 + indent * 6)
            add_formatted(p, m.group(2))
            i += 1
            continue

        # 空行
        if stripped == "":
            i += 1
            continue

        # 一般段落
        p = doc.add_paragraph()
        add_formatted(p, stripped)
        i += 1

    doc.save(dst)
    print(f"Saved: {dst}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
